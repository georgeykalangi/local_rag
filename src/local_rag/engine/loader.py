from __future__ import annotations

import fnmatch
import os
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class LoadedDocument:
    text: str
    source: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def load_document(path: Path | str) -> LoadedDocument:
    """Load a single document and extract its text content."""
    path = Path(path)
    suffix = path.suffix.lower()
    source = str(path)

    if suffix in (".txt", ".rst"):
        text = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(text=text, source=source, file_type=suffix)

    if suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(text=text, source=source, file_type=suffix)

    if suffix in (".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".rs",
                   ".c", ".cpp", ".h", ".hpp", ".rb", ".sh", ".yaml", ".yml",
                   ".toml", ".json", ".css", ".html"):
        text = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(
            text=text, source=source, file_type=suffix,
            metadata={"language": suffix.lstrip(".")},
        )

    if suffix == ".pdf":
        return _load_pdf(path)

    if suffix == ".docx":
        return _load_docx(path)

    raise ValueError(f"Unsupported file type: {suffix}")


def _load_pdf(path: Path) -> LoadedDocument:
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    text = "\n".join(pages)
    return LoadedDocument(
        text=text,
        source=str(path),
        file_type=".pdf",
        metadata={"page_count": len(doc)},
    )


def _load_docx(path: Path) -> LoadedDocument:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    text = "\n".join(paragraphs)
    return LoadedDocument(text=text, source=str(path), file_type=".docx")


def discover_files(
    folder: str,
    extensions: list[str],
    recursive: bool = True,
    max_size_mb: int = 10,
    ignore_file: str | None = None,
) -> list[Path]:
    """Walk a folder and return paths matching the given extensions and constraints."""
    folder_path = Path(folder)
    max_bytes = max_size_mb * 1024 * 1024
    ignore_patterns = _load_ignore_patterns(folder_path, ignore_file)

    results = []
    if recursive:
        walker = folder_path.rglob("*")
    else:
        walker = folder_path.glob("*")

    for entry in walker:
        if not entry.is_file():
            continue
        if entry.suffix.lower() not in extensions:
            continue
        if entry.stat().st_size > max_bytes:
            continue
        if _is_ignored(entry, folder_path, ignore_patterns):
            continue
        results.append(entry)

    return sorted(results)


def _load_ignore_patterns(folder: Path, ignore_file: str | None) -> list[str]:
    if not ignore_file:
        return []
    ignore_path = folder / ignore_file
    if not ignore_path.exists():
        return []
    patterns = []
    for line in ignore_path.read_text().splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            patterns.append(line)
    return patterns


def _is_ignored(path: Path, root: Path, patterns: list[str]) -> bool:
    rel = path.relative_to(root)
    rel_str = str(rel)
    for pattern in patterns:
        if pattern.endswith("/"):
            dir_pattern = pattern.rstrip("/")
            for parent in rel.parents:
                if fnmatch.fnmatch(str(parent), dir_pattern):
                    return True
                if str(parent) == dir_pattern:
                    return True
        elif fnmatch.fnmatch(rel_str, pattern):
            return True
    return False
