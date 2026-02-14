import os
from pathlib import Path

import pytest

from local_rag.engine.loader import load_document, discover_files, LoadedDocument

FIXTURES = Path(__file__).parent / "fixtures"


def test_load_txt_file():
    doc = load_document(FIXTURES / "sample.txt")
    assert isinstance(doc, LoadedDocument)
    assert "sample text file" in doc.text
    assert doc.file_type == ".txt"
    assert doc.source == str(FIXTURES / "sample.txt")


def test_load_markdown_file():
    doc = load_document(FIXTURES / "sample.md")
    assert "Sample Document" in doc.text
    assert "Section One" in doc.text
    assert doc.file_type == ".md"


def test_load_python_file():
    doc = load_document(FIXTURES / "sample.py")
    assert "def hello" in doc.text
    assert doc.file_type == ".py"


def test_load_pdf_file():
    pdf_path = FIXTURES / "sample.pdf"
    if not pdf_path.exists():
        pytest.skip("sample.pdf not in fixtures")
    doc = load_document(pdf_path)
    assert len(doc.text) > 0
    assert doc.file_type == ".pdf"
    assert doc.metadata.get("page_count") is not None


def test_load_docx_file(tmp_path):
    from docx import Document
    docx_path = tmp_path / "test.docx"
    d = Document()
    d.add_paragraph("First paragraph of test document.")
    d.add_paragraph("Second paragraph with more content.")
    d.save(str(docx_path))

    doc = load_document(docx_path)
    assert "First paragraph" in doc.text
    assert doc.file_type == ".docx"


def test_load_unsupported_file(tmp_path):
    bad_file = tmp_path / "data.bin"
    bad_file.write_bytes(b"\x00\x01\x02")
    with pytest.raises(ValueError, match="Unsupported"):
        load_document(bad_file)


def test_load_empty_file(tmp_path):
    empty = tmp_path / "empty.txt"
    empty.write_text("")
    doc = load_document(empty)
    assert doc.text == ""


def test_discover_files_respects_extensions(tmp_path):
    (tmp_path / "a.txt").write_text("hello")
    (tmp_path / "b.py").write_text("code")
    (tmp_path / "c.jpg").write_bytes(b"\xff\xd8")
    files = discover_files(str(tmp_path), extensions=[".txt", ".py"], recursive=True)
    names = {f.name for f in files}
    assert names == {"a.txt", "b.py"}


def test_discover_files_respects_max_size(tmp_path):
    small = tmp_path / "small.txt"
    small.write_text("small")
    big = tmp_path / "big.txt"
    big.write_text("x" * (2 * 1024 * 1024))  # 2MB
    files = discover_files(
        str(tmp_path), extensions=[".txt"], recursive=True, max_size_mb=1
    )
    names = {f.name for f in files}
    assert names == {"small.txt"}


def test_discover_files_respects_ragignore(tmp_path):
    (tmp_path / "keep.txt").write_text("keep")
    sub = tmp_path / "ignored_dir"
    sub.mkdir()
    (sub / "skip.txt").write_text("skip")
    (tmp_path / ".ragignore").write_text("ignored_dir/\n")
    files = discover_files(str(tmp_path), extensions=[".txt"], recursive=True,
                           ignore_file=".ragignore")
    names = {f.name for f in files}
    assert "keep.txt" in names
    assert "skip.txt" not in names


def test_discover_files_recursive_false(tmp_path):
    (tmp_path / "top.txt").write_text("top")
    sub = tmp_path / "sub"
    sub.mkdir()
    (sub / "nested.txt").write_text("nested")
    files = discover_files(str(tmp_path), extensions=[".txt"], recursive=False)
    names = {f.name for f in files}
    assert names == {"top.txt"}
